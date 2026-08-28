"""DOCX generation from a validated VoiceTranscript (speaker + timestamp
segments), using python-docx.

Returns bytes instead of writing to a path (unlike docx_service.py, which
serves the document-conversion pipeline where output files are tracked and
cleaned up via the database): voice transcription never persists the audio
or its output to disk - see app/bot/handlers/voice.py's module docstring.
"""

from __future__ import annotations

import io

from docx import Document as _new_docx_document
from docx.shared import Pt

from app.schemas.transcript import VoiceTranscript

_NO_SPEECH_TEXT = "Товуш аниқланмади ёки унда тушунарли нутқ топилмади."


def generate_transcript_docx(transcript: VoiceTranscript) -> bytes:
    doc = _new_docx_document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Овозли хабар транскрипцияси", level=0)

    if not transcript.segments:
        doc.add_paragraph(_NO_SPEECH_TEXT)
    else:
        for segment in transcript.segments:
            header = doc.add_paragraph()
            run = header.add_run(f"[{segment.start_time}] {segment.speaker}")
            run.bold = True
            doc.add_paragraph(segment.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
