"""DOCX generation from a validated DocumentResult, using python-docx.

Renders app/schemas/extraction.py::DocumentResult.text_blocks in reading
order (the authoritative sequence - see that module's docstring for why the
derived flat arrays are not used for rendering order).
"""

from __future__ import annotations

from datetime import UTC, datetime

from docx import Document as _new_docx_document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.extraction import DocumentResult, TextBlock

_UNCERTAIN_MARK = " [UNCERTAIN]"


def _render_heading(doc: DocxDocument, block: TextBlock) -> None:
    level = min(max(block.level or 1, 1), 6)
    text = block.text or ""
    if block.uncertain:
        text += _UNCERTAIN_MARK
    doc.add_heading(text, level=level)


def _render_paragraph(doc: DocxDocument, block: TextBlock) -> None:
    text = block.text or ""
    if block.uncertain:
        text += _UNCERTAIN_MARK
    paragraph = doc.add_paragraph(text)
    if block.type in ("signature", "stamp") and paragraph.runs:
        paragraph.runs[0].italic = True


def _render_list(doc: DocxDocument, block: TextBlock) -> None:
    if not block.items:
        # type="list" but no items were actually filled in - fall back to
        # the block's own text rather than silently dropping it (see
        # TextBlock.type docstring for why this defensive path exists).
        _render_paragraph(doc, block)
        return
    style = "List Number" if block.ordered else "List Bullet"
    for item in block.items:
        text = item.text
        if item.uncertain:
            text += _UNCERTAIN_MARK
        doc.add_paragraph(text, style=style)


def _render_table(doc: DocxDocument, block: TextBlock) -> None:
    table_data = block.table
    if table_data is None:
        # type="table" but no table payload - same defensive fallback.
        _render_paragraph(doc, block)
        return
    if table_data.title:
        title_text = table_data.title + (_UNCERTAIN_MARK if table_data.uncertain else "")
        title_paragraph = doc.add_paragraph(title_text)
        if title_paragraph.runs:
            title_paragraph.runs[0].bold = True

    columns = len(table_data.headers) or (len(table_data.rows[0]) if table_data.rows else 1)
    if columns == 0:
        return

    table = doc.add_table(rows=0, cols=columns)
    table.style = "Light Grid Accent 1"

    if table_data.headers:
        header_row = table.add_row().cells
        for i, header in enumerate(table_data.headers):
            header_row[i].text = header
            for paragraph in header_row[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    for row in table_data.rows:
        cells = table.add_row().cells
        for i, value in enumerate(row[:columns]):
            cells[i].text = value


_RENDERERS = {
    "heading": _render_heading,
    "paragraph": _render_paragraph,
    "signature": _render_paragraph,
    "stamp": _render_paragraph,
    "other": _render_paragraph,
    "list": _render_list,
    "table": _render_table,
}


def generate_docx(
    result: DocumentResult,
    source_filename: str,
    output_path: str,
) -> None:
    doc = _new_docx_document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    if result.metadata.title:
        title = doc.add_heading(result.metadata.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in result.text_blocks:
        renderer = _RENDERERS.get(block.type)
        if renderer:
            renderer(doc, block)

    if not result.text_blocks:
        doc.add_paragraph("No content could be extracted from this document.")

    _add_metadata_section(doc, result, source_filename)

    doc.save(output_path)


def _add_metadata_section(doc: DocxDocument, result: DocumentResult, source_filename: str) -> None:
    doc.add_page_break()
    doc.add_heading("Processing Information", level=2)
    doc.add_paragraph(f"Source file: {source_filename}")
    doc.add_paragraph(f"Processed: {datetime.now(UTC).isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Detected language: {result.language}")
    doc.add_paragraph(f"Detected document type: {result.document_type}")
    doc.add_paragraph(f"OCR/extraction confidence: {result.confidence:.2f}")
    if result.warnings:
        doc.add_paragraph("Warnings:")
        for warning in result.warnings:
            doc.add_paragraph(warning, style="List Bullet")
