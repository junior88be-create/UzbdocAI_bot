"""Tests for voice-transcript DOCX generation."""

from __future__ import annotations

import io

from docx import Document as DocxDocument

from app.schemas.transcript import TranscriptSegment, VoiceTranscript
from app.services.transcript_docx_service import generate_transcript_docx


def test_generate_transcript_docx_includes_speaker_and_timestamp():
    transcript = VoiceTranscript(
        language="uz",
        segments=[
            TranscriptSegment(speaker="Спикер 1", start_time="00:00", text="Ассалому алайкум."),
            TranscriptSegment(speaker="Спикер 2", start_time="00:12", text="Ваалайкум ассалом."),
        ],
    )

    docx_bytes = generate_transcript_docx(transcript)

    doc = DocxDocument(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "[00:00] Спикер 1" in all_text
    assert "Ассалому алайкум." in all_text
    assert "[00:12] Спикер 2" in all_text
    assert "Ваалайкум ассалом." in all_text


def test_generate_transcript_docx_reports_no_speech_when_empty():
    transcript = VoiceTranscript(language="unknown", segments=[])

    docx_bytes = generate_transcript_docx(transcript)

    doc = DocxDocument(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Товуш аниқланмади" in all_text


def test_generate_transcript_docx_bolds_the_speaker_header():
    transcript = VoiceTranscript(
        language="uz",
        segments=[TranscriptSegment(speaker="Спикер 1", start_time="00:00", text="Salom.")],
    )

    docx_bytes = generate_transcript_docx(transcript)

    doc = DocxDocument(io.BytesIO(docx_bytes))
    header_paragraph = next(p for p in doc.paragraphs if "Спикер 1" in p.text)
    assert header_paragraph.runs[0].bold is True
