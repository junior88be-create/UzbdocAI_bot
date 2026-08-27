"""Tests for DOCX generation."""

from __future__ import annotations

from docx import Document as DocxDocument

from app.schemas.extraction import DocumentResult, ListItem, TableBlock, TextBlock
from app.services.docx_service import generate_docx


def _sample_result() -> DocumentResult:
    return DocumentResult(
        document_type="official_letter",
        language="en",
        confidence=0.93,
        pages=1,
        text_blocks=[
            TextBlock(type="heading", text="Notice of Meeting", level=1, source_page=1),
            TextBlock(type="paragraph", text="Dear participants, please attend.", source_page=1),
            TextBlock(
                type="list",
                ordered=True,
                items=[ListItem(text="Agenda item one"), ListItem(text="Agenda item two")],
                source_page=1,
            ),
            TextBlock(
                type="table",
                table=TableBlock(title="Attendees", headers=["Name", "Role"], rows=[["Alice", "Chair"]]),
                source_page=1,
            ),
            TextBlock(type="paragraph", text="Illegible closing remark", source_page=1, uncertain=True),
        ],
        warnings=["Page 1 signature block partially illegible."],
    )


def test_generate_docx_creates_readable_file(tmp_path):
    result = _sample_result()
    output_path = tmp_path / "output.docx"

    generate_docx(result, source_filename="meeting_notice.pdf", output_path=str(output_path))

    assert output_path.exists()
    doc = DocxDocument(str(output_path))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Notice of Meeting" in all_text
    assert "Dear participants, please attend." in all_text
    assert "Agenda item one" in all_text
    assert "[UNCERTAIN]" in all_text
    assert "meeting_notice.pdf" in all_text

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Name"
    assert table.rows[1].cells[0].text == "Alice"


def test_generate_docx_handles_empty_document(tmp_path):
    result = DocumentResult(document_type="unknown", language="unknown", pages=1)
    output_path = tmp_path / "empty.docx"

    generate_docx(result, source_filename="blank.pdf", output_path=str(output_path))

    assert output_path.exists()
    doc = DocxDocument(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "No content could be extracted" in all_text


def test_generate_docx_falls_back_to_text_for_table_type_with_no_table_payload(tmp_path):
    """Regression: a block declared type="table" but with `table` left
    unset used to be silently dropped entirely - see TextBlock.type
    docstring. It must now render its raw text instead of vanishing."""
    result = DocumentResult(
        document_type="notice",
        language="uz",
        pages=1,
        text_blocks=[
            TextBlock(type="heading", text="Notice", level=1, source_page=1),
            TextBlock(type="table", text="Ҳисоб рақами: 6000133137", table=None, source_page=1),
        ],
    )
    output_path = tmp_path / "fallback_table.docx"

    generate_docx(result, source_filename="notice.pdf", output_path=str(output_path))

    doc = DocxDocument(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Ҳисоб рақами: 6000133137" in all_text


def test_generate_docx_falls_back_to_text_for_list_type_with_no_items(tmp_path):
    result = DocumentResult(
        document_type="notice",
        language="uz",
        pages=1,
        text_blocks=[
            TextBlock(type="list", text="Кечиктирилган ҳар куни учун 0,1%.", items=[], source_page=1),
        ],
    )
    output_path = tmp_path / "fallback_list.docx"

    generate_docx(result, source_filename="notice.pdf", output_path=str(output_path))

    doc = DocxDocument(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Кечиктирилган ҳар куни учун 0,1%." in all_text
